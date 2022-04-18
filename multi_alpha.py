import docx
import os
def queue_text(seq, useq):
    if seq == 1:
        for root, dirs, files in os.walk('D:/un/diss/project/test_text/1'):
            for file in files:
                doc = docx.Document(r"D:/un/diss/project/test_text/1/"+file)
                y = 0
                txt5 = ''
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        txt = run.text
                        txt2=txt.lower()
                        txt3=txt2.replace(",", "")
                        txt4 = txt3.replace(".", "")
                        txt5=txt4.strip()
                        print(txt5)
                        print(y)
                        if run.bold:
                            print('1')
                            a = str(run.text)
                            g = "'"+a+"'"
                            query1 = "MERGE(f:Fish{name:"+g+"})"
                            conn.query(query1, db='DB1')
                        if run.italic:
                            print('2')
                            q = str(run.text)
                            t = "'"+q+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) MERGE(d:Dome{name:"+t+"}) MERGE (f)-[:isa]->(d)"
                            conn.query(query1, db='DB1')
                        if txt5 in ('усов', 'усы'):
                            y = 1
                            print('a')
                            txt5 = ''
                        if run.underline and y == 1:
                            print('3')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.barbel="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 =='губы':
                            y = 2
                            print('b')
                            txt5 = ''
                        if run.underline and y == 2:
                            print('4')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.lips="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 == 'размер':
                            y = 3
                            print('c')
                            txt5 = ''
                        if run.underline and y == 3:
                            print('5')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.size="+h
                            conn.query(query1, db='DB1')
                            y = 0
    if seq == 2:
        for root, dirs, files in os.walk('D:/un/diss/project/test_text/2'):
            for file in files:
                doc = docx.Document(r"D:/un/diss/project/test_text/2/"+file)
                y = 0
                txt5 = ''
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        txt = run.text
                        txt2=txt.lower()
                        txt3=txt2.replace(",", "")
                        txt4 = txt3.replace(".", "")
                        txt5=txt4.strip()
                        print(txt5)
                        print(y)
                        if run.bold:
                            print('1')
                            a = str(run.text)
                            g = "'"+a+"'"
                            query1 = "MERGE(f:Fish{name:"+g+"})"
                            conn.query(query1, db='DB1')
                        if run.italic:
                            print('2')
                            q = str(run.text)
                            t = "'"+q+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) MERGE(d:Dome{name:"+t+"}) MERGE (f)-[:isa]->(d)"
                            conn.query(query1, db='DB1')
                        if txt5 in ('усов', 'усы'):
                            y = 1
                            print('a')
                            txt5 = ''
                        if run.underline and y == 1:
                            print('3')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.barbel="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 =='губы':
                            y = 2
                            print('b')
                            txt5 = ''
                        if run.underline and y == 2:
                            print('4')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.lips="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 == 'размер':
                            y = 3
                            print('c')
                            txt5 = ''
                        if run.underline and y == 3:
                            print('5')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.size="+h
                            conn.query(query1, db='DB1')
                            y = 0
    if seq == 3:
        for root, dirs, files in os.walk('D:/un/diss/project/test_text/3'):
            for file in files:
                doc = docx.Document(r"D:/un/diss/project/test_text/3/"+file)
                y = 0
                txt5 = ''
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        txt = run.text
                        txt2=txt.lower()
                        txt3=txt2.replace(",", "")
                        txt4 = txt3.replace(".", "")
                        txt5=txt4.strip()
                        print(txt5)
                        print(y)
                        if run.bold:
                            print('1')
                            a = str(run.text)
                            g = "'"+a+"'"
                            query1 = "MERGE(f:Fish{name:"+g+"})"
                            conn.query(query1, db='DB1')
                        if run.italic:
                            print('2')
                            q = str(run.text)
                            t = "'"+q+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) MERGE(d:Dome{name:"+t+"}) MERGE (f)-[:isa]->(d)"
                            conn.query(query1, db='DB1')
                        if txt5 in ('усов', 'усы'):
                            y = 1
                            print('a')
                            txt5 = ''
                        if run.underline and y == 1:
                            print('3')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.barbel="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 =='губы':
                            y = 2
                            print('b')
                            txt5 = ''
                        if run.underline and y == 2:
                            print('4')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.lips="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 == 'размер':
                            y = 3
                            print('c')
                            txt5 = ''
                        if run.underline and y == 3:
                            print('5')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.size="+h
                            conn.query(query1, db='DB1')
                            y = 0
    if seq == 4:
        for root, dirs, files in os.walk('D:/un/diss/project/test_text/4'):
            for file in files:
                doc = docx.Document(r"D:/un/diss/project/test_text/4/"+file)
                y = 0
                txt5 = ''
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        txt = run.text
                        txt2=txt.lower()
                        txt3=txt2.replace(",", "")
                        txt4 = txt3.replace(".", "")
                        txt5=txt4.strip()
                        print(txt5)
                        print(y)
                        if run.bold:
                            print('1')
                            a = str(run.text)
                            g = "'"+a+"'"
                            query1 = "MERGE(f:Fish{name:"+g+"})"
                            conn.query(query1, db='DB1')
                        if run.italic:
                            print('2')
                            q = str(run.text)
                            t = "'"+q+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) MERGE(d:Dome{name:"+t+"}) MERGE (f)-[:isa]->(d)"
                            conn.query(query1, db='DB1')
                        if txt5 in ('усов', 'усы'):
                            y = 1
                            print('a')
                            txt5 = ''
                        if run.underline and y == 1:
                            print('3')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.barbel="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 =='губы':
                            y = 2
                            print('b')
                            txt5 = ''
                        if run.underline and y == 2:
                            print('4')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.lips="+h
                            conn.query(query1, db='DB1')
                            y = 0
                        if txt5 == 'размер':
                            y = 3
                            print('c')
                            txt5 = ''
                        if run.underline and y == 3:
                            print('5')
                            с = str(run.text)
                            h = "'"+с+"'"
                            query1 = "MATCH(f:Fish{name:"+g+"}) SET f.size="+h
                            conn.query(query1, db='DB1')
                            y = 0
    print('done')
    return seq

import threading
import os
import time
if __name__ == '__main__':
    start_time = time.clock()
    p1 = threading.Thread(target = queue_text, args =(1, '' ) ) 
    p1.start()
    p2 = threading.Thread(target = queue_text, args =(2, '' ) ) 
    p2.start()
    p3 = threading.Thread(target = queue_text, args =(3, '' ) ) 
    p3.start()
    p4 = threading.Thread(target = queue_text, args =(4, '' ) ) 
    p4.start()

